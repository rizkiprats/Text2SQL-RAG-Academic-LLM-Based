import psycopg2
from langchain_ollama import ChatOllama
from langchain_chroma import Chroma as LangchainChroma
from langchain_huggingface.embeddings import HuggingFaceEmbeddings
from langchain.schema.runnable import RunnableMap, RunnablePassthrough, RunnableLambda
from langchain.schema import Document as SchemaDocument
from cache_manager import SQLResponseCache
import json
import re
import datetime

import os
from langchain_core.documents import Document as LangchainDocument
from langchain_community.vectorstores import Chroma as ChromaVectorStore

from constant import (
    MODELS,
    DB_CONFIG,
)

from langchain_core.messages import AIMessage

from prompt import prompt_sql_generator, prompt_classify_question 
from prompt import prompt_related_question_check, prompt_summary_question

from utils import execute_query_and_return_df, dataframe_to_json, json_to_dataframe

class TextRAG:
    def __init__(self):
        self.embeddings = HuggingFaceEmbeddings(
            model_name=MODELS["embedding_model"], encode_kwargs={"normalize_embeddings": True})

        # Gunakan Ollama lokal
        self.model = ChatOllama(
            model=MODELS["general_llm_model"], base_url="http://localhost:11434", temperature=0.7, num_ctx=6114)

        self.retriever, self.docs = self._indexing_vectore()

    def load_all_documents(self, folder_path):
        all_contents = []
        for filename in os.listdir(folder_path):
            path = os.path.join(folder_path, filename)
            try:
                if filename.endswith(".pdf"):
                    from PyPDF2 import PdfReader
                    reader = PdfReader(path)
                    text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
                    document = LangchainDocument(page_content=text, metadata={"document file": filename})
                    all_contents.append(document)

                elif filename.endswith(".docx"):
                    from docx import Document as DocumentDocx
                    doc = DocumentDocx(path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    document = LangchainDocument(page_content=text, metadata={"document file": filename})
                    all_contents.append(document)

                elif filename.endswith(".xlsx") or filename.endswith(".xls"):
                    import pandas as pd
                    df = pd.read_excel(path)
                    text = df.to_string(index=False)
                    document = LangchainDocument(page_content=text, metadata={"document file": filename})
                    all_contents.append(document)

            except Exception as e:
                print(f"Gagal membaca file {filename}: {e}")
        
        return all_contents

    def _indexing_vectore(self):
        # Load documents and create vector store
        documents = self.load_all_documents("document_files")

        if documents:
            doc_vectorstore = ChromaVectorStore.from_documents(
                documents, collection_name="documents_docs", embedding=self.embeddings)
            doc_retriever = doc_vectorstore.as_retriever(
                search_kwargs={"k": 1})

            return doc_retriever, doc_vectorstore
        else:
            return None, None

    def run_related_context_check(self, question: str, last_response: str, last_data_json: str) -> bool:
        prompt = prompt_related_question_check

        # fungsi untuk debug payload ke invoke_pipeline
        def debug_print(x):
            print("\n--- DEBUG: payload ke invoke_pipeline ---")
            import pprint
            pprint.pprint(x)
            return x
        
        # Define RAG pipeline
        rag_chain = (
            RunnableMap(
                {"last_response": last_response, "last_data_json": last_data_json, "question": RunnablePassthrough()})
            | prompt
            | RunnableLambda(debug_print)
            | self.model
        )

        # Run pipeline
        try:
            response = rag_chain.invoke(question)
            print("response question related check: ", response)
        except Exception as e:
            print(f"Error during question related check: {e}")
            response = AIMessage(content=f"Error during question related check: {e}")

        # Extract response content
        if isinstance(response, AIMessage):
            response_content = response.content
        else:
            response_content = str(response)

        if "true" in response_content.lower():
            return True
        
        return False

    def run_summary_context(self, question: str, last_response: str, data_json: str) -> str:
        prompt = prompt_summary_question

        # fungsi untuk debug payload ke invoke_pipeline
        def debug_print(x):
            print("\n--- DEBUG: payload ke invoke_pipeline ---")
            import pprint
            pprint.pprint(x)
            return x
        
        # Define RAG pipeline
        rag_chain = (
            RunnableMap(
                {"last_response": last_response, "data_json": data_json, "question": RunnablePassthrough()})
            | prompt
            | RunnableLambda(debug_print)
            | self.model
        )

        # Run pipeline
        try:
            response = rag_chain.invoke(question)
            print("response summary context : ", response)
        except Exception as e:
            print(f"Error during summary context: {e}")
            response = AIMessage(content=f"Error during summary context: {e}")

        # Extract response content
        if isinstance(response, AIMessage):
            response_content = response.content
        else:
            response_content = str(response)

        return response_content
    

class Text2SQLRAG:
    def __init__(self, db_config: dict = DB_CONFIG):
        """
        A class for generating SQL queries based on natural language text.
        """
        self.db_config = db_config

        self.embeddings = HuggingFaceEmbeddings(
            model_name=MODELS["embedding_model_sql"], encode_kwargs={"normalize_embeddings": True})

        self.retriever, self.docs = self._indexing_vectore()
        self.retriever_analysis, self.docs_analysis = self._indexing_vectore_analysis()


        # Gunakan Ollama lokal
        self.model = ChatOllama(
            model=MODELS["sql_llm_model"], base_url="http://localhost:11434", temperature=0.1, num_ctx=6114)
        self.model_answer = ChatOllama(
            model=MODELS["general_llm_model"], base_url="http://localhost:11434", temperature=0.7, num_ctx=6114)

        self.cache = SQLResponseCache()

    def _fetch_postgres_schema(self):
        """
        Mengambil skema database PostgreSQL secara langsung.
        """

        schema_texts = []
        try:
            print("Start fetching schema from PostgreSQL Database...")

            conn = psycopg2.connect(**self.db_config)
            cur = conn.cursor()

            # Ambil Schema Database
            cur.execute("""
                    SELECT table_name FROM information_schema.tables 
                    WHERE table_schema = 'public'
            """)

            tables = cur.fetchall()
            print("database tables : ", tables)

            for (table_name,) in tables:
                # Ambil skema dari setiap tabel
                cur.execute(
                    f"SELECT column_name, data_type FROM information_schema.columns WHERE table_name = '{table_name}'")
                columns = cur.fetchall()

                column_list = ",\n".join(
                    f"{col} ({dtype.upper()})" for col, dtype in columns)
                schema_text = f"Table: {table_name}\nColumns:\n{column_list}\n\n"

                schema_texts.append(SchemaDocument(page_content=schema_text, metadata={"table": table_name}))

            cur.close()
            conn.close()

            print("Finished fetching schema from PostgreSQL Database.")
            print("schema_texts", schema_texts)

        except Exception as e:
            print(f"Error fetching schema: {e}")

        return schema_texts

    def _indexing_vectore(self):
        """
        Mengindeks skema PostgreSQL ke dalam vektor untuk pencarian.
        """
        
        # Ambil skema database dari PostgreSQL
        schema_texts = self._fetch_postgres_schema()

        if schema_texts:
            print("Start indexing schema from PostgreSQL Database...")

            # Simpan embeddings
            docs = schema_texts

            vectorstore = LangchainChroma.from_documents(docs, collection_name="sql_docs", embedding=self.embeddings)
            retriever = vectorstore.as_retriever(search_kwargs={"k": 4})

            print('Finished indexing schema from PostgreSQL Database.')
            schema_indexes = "\n".join([doc.page_content for doc in docs])

            print("sql retriever :", retriever)
            print("sql docs :", docs)
            print("schema_indexes :", schema_indexes)

            return retriever, docs
        else:
            print('No schema database found')
            return None, None

    def _fetch_postgres_schema_analysis(self):
        """
        Mengambil skema database PostgreSQL secara langsung untuk analisis skema dan klasifikasi pertanyaan.
        """

        schema_texts = []
        schema_metadata = {}
        try:
            print("Start fetching schema from PostgreSQL Database for analysis...")

            conn = psycopg2.connect(**self.db_config)
            cur = conn.cursor()

            # Ambil daftar tabel
            cur.execute("""
                    SELECT table_name FROM information_schema.tables 
                    WHERE table_schema = 'public'
            """)
            tables = cur.fetchall()
            print("database tables : ", tables)

            for (table_name,) in tables:
                # Ambil kolom dan tipe data dari setiap tabel
                cur.execute(f"""
                    SELECT column_name, data_type 
                    FROM information_schema.columns 
                    WHERE table_name = '{table_name}'
                """)
                columns = cur.fetchall()

                # Ambil foreign key untuk tabel
                cur.execute(f"""
                    SELECT
                        kcu.column_name,
                        ccu.table_name AS foreign_table_name,
                        ccu.column_name AS foreign_column_name
                    FROM 
                        information_schema.key_column_usage AS kcu
                    JOIN 
                        information_schema.constraint_column_usage AS ccu
                    ON 
                        kcu.constraint_name = ccu.constraint_name
                    WHERE 
                        kcu.table_name = '{table_name}'
                """)
                foreign_keys = cur.fetchall()

                # Format metadata tabel
                column_list = ",\n".join(
                    f"{col} ({dtype.upper()})" for col, dtype in columns)
                foreign_key_list = "\n".join(
                    f"{col} -> {foreign_table}({foreign_col})"
                    for col, foreign_table, foreign_col in foreign_keys
                )
                schema_text = f"Table: {table_name}\nColumns:\n{column_list}\n"
                if foreign_keys:
                    schema_text += f"Foreign Keys:\n{foreign_key_list}\n"
                schema_text += "\n"

                schema_texts.append(SchemaDocument(page_content=schema_text, metadata={"table": table_name}))
                schema_metadata[table_name] = {
                    "columns": columns,
                    "foreign_keys": foreign_keys
                }

            cur.close()
            conn.close()

            print("Finished fetching schema from PostgreSQL Database for analysis.")
            print("schema_texts for analysis", schema_texts)
            print("schema_metadata for analysis", schema_metadata)

        except Exception as e:
            print(f"Error fetching schema for analysis: {e}")
        
        return schema_texts, schema_metadata

    def _indexing_vectore_analysis(self):
        """
        Mengindeks skema PostgreSQL ke dalam vektor untuk pencarian.
        """

        # Ambil skema database dari PostgreSQL
        schema_texts, _ = self._fetch_postgres_schema_analysis()

        if schema_texts:
            print("Start indexing schema from PostgreSQL Database for analysis...")

            # Simpan embeddings
            docs = schema_texts

            vectorstore = LangchainChroma.from_documents(docs, collection_name="sql_analysis_docs", embedding=self.embeddings)
            retriever = vectorstore.as_retriever(search_kwargs={"k": 4})

            print('Finished indexing schema from PostgreSQL Database for analysis.')
            schema_indexes = "\n".join([doc.page_content for doc in docs])

            print("sql retriever for analysis :", retriever)
            print("sql docs for analysis :", docs)
            print("schema indexes for analysis:", schema_indexes)
            
            return retriever, docs
        else:
            print('No schema database found for analysis')
            
            return None, None
        
    def run_sql_rag(self, question: str):
        if self.retriever:
            docs = self.retriever.invoke(question)
            schema_retrieve = "\n".join([doc.page_content for doc in docs])
            context = schema_retrieve
        else:
            context = ""

        def context_fn(_): return context

        current_date = datetime.datetime.now()
        current_date_str = current_date.strftime("%Y-%m-%d")
        print("current_date_str : ", current_date_str)
        def current_date(_): return current_date_str

        print("The model:", self.model)

        prompt = prompt_sql_generator

        # fungsi untuk debug payload ke invoke_pipeline
        def debug_print(x):
            print("\n--- DEBUG: payload ke invoke_pipeline ---")
            import pprint
            pprint.pprint(x)
            return x
        
        # Define RAG pipeline
        rag_chain = (
            RunnableMap(
                {"context": context_fn, "date": current_date, "question": RunnablePassthrough()})
            | prompt
            | RunnableLambda(debug_print)
            | self.model
        )

        # Run pipeline
        try:
            response = rag_chain.invoke(question)
            print("response question rag: ", response)
        except Exception as e:
            print(f"Error during SQL RAG pipeline execution: {e}")
            response = AIMessage(content=f"Error during SQL RAG pipeline execution: {e}")
    
        # Extract response content sql generate
        if isinstance(response, AIMessage):
            response_content = response.content
        else:
            response_content = str(response)

        # Extract SQL query from response content sql generate
        sql_match = re.search(
            r"```sql\s+(SELECT[\s\S]*?)```", response_content, re.IGNORECASE)
        if not sql_match:
            sql_match = re.search(
                r"(SELECT[\s\S]*?;)", response_content, re.IGNORECASE)
        if not sql_match:
            sql_match = re.search(
                r"(SELECT[\s\S]*)", response_content, re.IGNORECASE)

        sql_query = sql_match.group(1).strip() if sql_match else None
        reasoning = response_content

        print("response : ", response_content)
        print("sql_query : ", sql_query)
        print("reasoning : ", reasoning)

        # Execute SQL query if available
        if sql_query:
            try:
                df, error = execute_query_and_return_df(sql_query, return_error=True)
            except Exception as e:
                print(f"Error executing SQL query: {str(e)}")
                df = None
                error = str(e)
        else:
            df = None
            error = "No SQL query found in the response."
        
        return df, error

    def run_classify_question(self, question: str):

        if self.retriever_analysis:
            docs = self.retriever_analysis.invoke(question)
            schema_retrieve = "\n".join([doc.page_content for doc in docs])
            context = schema_retrieve
        else:
            context = ""

        print("Context for classification:", context)
        def context_fn(_): return context

        prompt = prompt_classify_question

        format_output = "{\"queryType\": \"GENERAL_QUESTION\" | \"DATA_QUESTION\" | \"OUT_OF_SCOPE\"}"

        # fungsi untuk debug payload ke invoke_pipeline
        def debug_print(x):
            print("\n--- DEBUG: payload ke invoke_pipeline ---")
            import pprint
            pprint.pprint(x)
            return x
        
        # Define RAG pipeline
        rag_chain = (
            RunnableMap(
                {"context": context_fn, "format_output": format_output, "question": RunnablePassthrough()})
            | prompt
            | RunnableLambda(debug_print)
            | self.model
        )

        try:
            # Run pipeline
            response = rag_chain.invoke(question)
            print("response analysis classification question: ", response)
        except Exception as e:
            print(f"Error during classify question: {e}")
            response = AIMessage(content=json.dumps({
                "queryType": "GENERAL_QUESTION",
                "message": str(e)
            }))
        
        # Extract response content classification question
        if isinstance(response, AIMessage):
            response_content = response.content
        else:
            response_content = str(response)

        # Extract Json from response content classify question
        json_match = re.search(r"```json\s+(\{.*?\"queryType\".*?\})```", response_content, re.IGNORECASE)
        if not json_match:
            json_match = re.search(
                r"(\{.*?\"queryType\".*?\})", response_content, re.IGNORECASE)
        if not json_match:
            json_match = re.search(
                r"(\{.*?\"queryType\".*?\})", response_content, re.IGNORECASE)
            
        json_match = json_match.group(1).strip() if json_match else None
        reasoning = response_content

        print("response : ", response_content)
        print("json : ", json_match)
        print("reasoning : ", reasoning)

        return json_match